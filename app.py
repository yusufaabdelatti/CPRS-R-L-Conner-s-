import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, smtplib, re
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

# PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, Image as RLImage, HRFlowable,
                                 PageBreak, KeepTogether)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ══════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════
GMAIL_USER      = "Wijdan.psyc@gmail.com"
GMAIL_PASS      = "rias eeul lyuu stce"
RECIPIENT_EMAIL = "Wijdan.psyc@gmail.com"
LOGO_FILE       = "logo.png"

CLINIC_BLUE_RGB = RGBColor(0x8B, 0x73, 0x55)
DARK_RGB        = RGBColor(0x1C, 0x19, 0x17)
WARM_RGB        = RGBColor(0x8B, 0x73, 0x55)
ACCENT_RGB      = RGBColor(0xC4, 0x95, 0x6A)

# ══════════════════════════════════════════════════════════════
#  80 ITEMS — English + Arabic
# ══════════════════════════════════════════════════════════════
ITEMS_EN = [
    "Angry and resentful",
    "Difficulty doing or completing homework",
    "Is always 'on the go' or acts as if driven by a motor",
    "Timid, easily frightened",
    "Everything must be just so",
    "Has no friends",
    "Stomach aches",
    "Fights",
    "Avoids or has difficulties engaging in tasks that require sustained mental effort",
    "Has difficulty sustaining attention in tasks or play activities",
    "Argues with adults",
    "Fails to complete assignments",
    "Hard to control in malls or while grocery shopping",
    "Afraid of people",
    "Keeps checking things over and over again",
    "Loses friends quickly",
    "Aches and Pains",
    "Restless or overactive",
    "Has trouble concentrating in class",
    "Does not seem to listen to what is being said",
    "Loses temper",
    "Needs close supervision to get through assignments",
    "Runs about or climbs excessively in inappropriate situations",
    "Afraid of new situations",
    "Fussy about cleanliness",
    "Does not know how to make friends",
    "Gets aches and pains or stomachaches before school",
    "Excitable, impulsive",
    "Does not follow through on instructions and fails to finish schoolwork",
    "Has difficulty organising tasks and activities",
    "Irritable",
    "Restless in the 'squirmy sense'",
    "Afraid of being alone",
    "Things must be done the same way every time",
    "Does not get invited over to friends' houses",
    "Headaches",
    "Fails to finish things he/she starts",
    "Inattentive, easily distracted",
    "Talks excessively",
    "Actively defies or refuses to comply with adults' requests",
    "Fails to give close attention to detail or makes careless mistakes",
    "Has difficulty waiting in lines or awaiting turn in group situations",
    "Has a lot of fears",
    "Has rituals that he/she must go through",
    "Distractibility or attention span problem",
    "Complains about being sick even when nothing is wrong",
    "Temper outbursts",
    "Gets distracted when given instructions to do something",
    "Interrupts or intrudes on others (butts into conversations or games)",
    "Forgetful in daily activities",
    "Cannot grasp arithmetic",
    "Will run around between mouthfuls at meals",
    "Afraid of the dark, animals or bugs",
    "Sets very high goals for self",
    "Fidgets with hands or feet or squirms in seat",
    "Short attention span",
    "Touchy or easily annoyed by others",
    "Has sloppy handwriting",
    "Has difficulty playing or engaging in leisure activities quietly",
    "Shy, withdrawn",
    "Blames others for his/her mistakes or misbehaviour",
    "Fidgeting",
    "Messy or disorganised at home or school",
    "Gets upset if someone rearranges his/her things",
    "Clings to parents or other adults",
    "Disturbs other children",
    "Deliberately does things that annoy other people",
    "Demands must be met immediately — easily frustrated",
    "Only attends if it is something he/she is very interested in",
    "Spiteful or vindictive",
    "Loses things necessary for tasks (pencils, books, tools or toys)",
    "Feels inferior to others",
    "Seems tired or slowed down at times",
    "Spelling is poor",
    "Cries often and easily",
    "Leaves seat in classroom or where remaining seated is expected",
    "Mood changes quickly and drastically",
    "Easily frustrated efforts",
    "Easily distracted by extraneous stimuli",
    "Blurts out answers before questions have been completed",
]

ITEMS_AR = [
    "مستاء وغاضب",
    "يعاني من صعوبة في أداء الواجب أو إنهاءه",
    "دائما يريد الحركة أو يتصرف كأنه مدفوع بموتور",
    "خجول ـ يخاف بسهولة",
    "كل شيء يجب أن يكون دقيقاً ومضبوطاً",
    "ليس لديه أصدقاء",
    "يعاني من أمراض المعدة",
    "يتخاتق ويتشاجر",
    "يتجنب أو لديه صعوبة في عمل شيء يحتاج إلى تركيز ذهني (واجب المدرسة)",
    "يعاني من صعوبة في التركيز فترة طويلة في الأعمال أو اللعب",
    "يجادل مع الكبار",
    "يفشل في إنهاء مهماته أو واجباته",
    "صعب السيطرة عليه في الأسواق التجارية أو أثناء شراء احتياجات المنزل",
    "يخاف من الناس",
    "يتأكد من الأشياء مراراً وتكراراً",
    "يخسر أصحابه بسرعة",
    "عنده أوجاع وآلام",
    "لا يهدأ وكثير النشاط والحركة غير مستقر",
    "يعاني من مشاكل في التركيز في الفصل",
    "لا يستمع لما يقال إليه",
    "يفقد أعصابه",
    "يحتاج إلى إشراف دائم لينتهي من واجباته",
    "يجري أو يتسلق كثيراً في موقف لا يصح فيه هذا التصرف",
    "يخاف من المواقف الجديدة",
    "يهتم بالنظافة إلى حد مزعج أو كبير",
    "لا يعرف كيف يعمل صداقات",
    "يعاني من أوجاع وآلام أو ألم بالمعدة قبل الذهاب للمدرسة",
    "سهل الاستثارة ومندفع",
    "لا يتبع التعليمات ويفشل في إنهاء واجباته في العمل أو الدراسة",
    "يعاني من صعوبة في تنظيم الواجبات والنشاطات",
    "متهيج",
    "كثير الحركة أو قلق",
    "يخاف من البقاء بمفرده",
    "لابد من عمل الأشياء بنفس الطريقة كل مرة",
    "لا يدعوه أحد من أصدقائه لزيارته بمنزله",
    "يعاني من الصداع",
    "يفشل في إنهاء الأشياء التي بدأها",
    "قليل التركيز، سهل أن تتشتت تركيزه",
    "يتكلم كثيراً",
    "يعاند أو يرفض بقوة أن يلتزم بطلبات الكبار",
    "يفشل أن يعطي انتباهه للتفاصيل ويرتكب أخطاء في المدرسة أو العمل أو أي نشاط آخر",
    "يعاني من صعوبة في الانتظار في الطابور أو انتظار دوره في اللعب أو المواقف الجماعية",
    "يعاني من مخاوف كثيرة",
    "لديه طقوس لابد أن يؤديها",
    "تشتت الانتباه أو مشكلة في مدة الانتباه",
    "يشكو من المرض حتى عندما لا يكون هناك شيء خاطئ",
    "نوبات غضب",
    "ينشغل عند إعطائه تعليمات للقيام بشيء ما",
    "يقاطع الآخرين أو يتدخل فيهم (يقتحم المحادثات أو الألعاب)",
    "ناسٍ في أنشطة حياته اليومية",
    "لا يستطيع فهم الحساب",
    "يركض بين قضمات الطعام أثناء الوجبات",
    "يخاف من الظلام أو الحيوانات أو الحشرات",
    "يضع لنفسه أهدافاً عالية جداً",
    "يعبث بيديه أو قدميه أو يتلوى في مقعده",
    "مدة الانتباه قصيرة",
    "حساس أو يتضايق بسهولة من الآخرين",
    "خطه سيء",
    "يعاني من صعوبة في اللعب أو الانخراط في أنشطة وقت الفراغ بهدوء",
    "خجول ومنسحب",
    "يلوم الآخرين على أخطائه أو سوء تصرفه",
    "كثير الحركة والتململ",
    "فوضوي أو غير منظم في المنزل أو المدرسة",
    "يضطرب إذا رتّب أحد أغراضه",
    "يتعلق بوالديه أو بالغين آخرين",
    "يزعج الأطفال الآخرين",
    "يتعمد فعل أشياء تزعج الآخرين",
    "يجب تلبية مطالبه فوراً — يُحبَط بسهولة",
    "ينتبه فقط إذا كان الشيء يثير اهتمامه جداً",
    "حقود أو انتقامي",
    "يفقد الأشياء الضرورية للمهام (أقلام، كتب، أدوات أو ألعاب)",
    "يشعر بأنه أقل من الآخرين",
    "يبدو متعباً أو بطيئاً في بعض الأحيان",
    "إملاؤه ضعيف",
    "يبكي كثيراً وبسهولة",
    "يترك مقعده في الفصل أو حيث يُتوقع منه الجلوس",
    "مزاجه يتغير بسرعة وبشكل كبير",
    "يُحبَط بسهولة في جهوده",
    "يتشتت بسهولة بمحفزات خارجية",
    "يبلغ بالإجابات قبل اكتمال الأسئلة",
]

# ══════════════════════════════════════════════════════════════
#  SUBSCALES
# ══════════════════════════════════════════════════════════════
SUBSCALES = {
    "A": {"name_en": "Oppositional",               "name_ar": "المعارضة",
          "items": [8,11,21,31,40,47,57,61,67,70], "color": "#C62828"},
    "B": {"name_en": "Cognitive/Inattention",      "name_ar": "الإدراك / قصور الانتباه",
          "items": [2,9,10,12,19,22,29,30,37,38,41,45,48,50,51,56,58,63,69,71,74], "color": "#1565C0"},
    "C": {"name_en": "Hyperactivity",              "name_ar": "فرط الحركة",
          "items": [3,13,18,23,28,32,39,42,49,52,55,59,62,66,76,80], "color": "#E65100"},
    "D": {"name_en": "Anxious-Shy",                "name_ar": "القلق والخجل",
          "items": [4,14,24,33,43,53,60,65,72], "color": "#6A1B9A"},
    "E": {"name_en": "Perfectionism",              "name_ar": "الكمالية",
          "items": [5,15,25,34,44,54,64], "color": "#00695C"},
    "F": {"name_en": "Social Problems",            "name_ar": "المشكلات الاجتماعية",
          "items": [6,16,26,35,72], "color": "#0277BD"},
    "G": {"name_en": "Psychosomatic",              "name_ar": "الأعراض النفسجسمية",
          "items": [7,17,27,36,46,73], "color": "#4E342E"},
    "H": {"name_en": "ADHD Index",                 "name_ar": "مؤشر ADHD",
          "items": [3,6,9,10,18,20,22,28,31,32,38,39,45,47,48,49,55,57,62,66,80], "color": "#283593"},
    "I": {"name_en": "CGI: Restless-Impulsive",   "name_ar": "CGI: الاندفاعية وعدم الهدوء",
          "items": [3,13,18,21,23,28,31,32,39,42,47,49,52,55,57,59,62,66,68,76,77,78,80], "color": "#AD1457"},
    "J": {"name_en": "CGI: Emotional Lability",   "name_ar": "CGI: التقلب الانفعالي",
          "items": [1,21,31,47,57,75,77,78], "color": "#558B2F"},
    "K": {"name_en": "CGI: Total",                 "name_ar": "المؤشر العام",
          "items": [3,6,9,10,12,13,18,19,20,21,22,28,31,32,38,39,42,45,47,48,49,55,57,62,75,77,79,80], "color": "#1B5E20"},
    "L": {"name_en": "DSM-IV: Inattentive",        "name_ar": "نقص الانتباه DSM-IV",
          "items": [9,10,12,20,29,38,41,45,50,56,71,79], "color": "#4527A0"},
    "M": {"name_en": "DSM-IV: Hyperactive-Impulsive", "name_ar": "فرط الحركة والاندفاعية DSM-IV",
          "items": [3,18,23,28,32,39,42,49,52,55,59,62,76,80], "color": "#BF360C"},
    "N": {"name_en": "DSM-IV: Total",              "name_ar": "مختلط DSM-IV",
          "items": [3,9,10,12,18,20,23,28,29,32,38,39,41,42,45,49,50,52,55,56,59,62,71,76,79,80], "color": "#33691E"},
}

NORMS = {
    "A": (6.8,4.2), "B": (13.5,7.8), "C": (10.2,6.5), "D": (5.8,4.0),
    "E": (5.2,3.8), "F": (3.5,2.8),  "G": (2.8,2.5),  "H": (11.0,6.8),
    "I": (13.5,7.5),"J": (5.5,3.9),  "K": (35.0,16.0),"L": (11.2,6.8),
    "M": (10.2,6.5),"N": (18.5,10.5),
}

def get_level_en(t):
    if t >= 70:   return "Markedly Atypical"
    elif t >= 65: return "Likely Concern"
    elif t >= 60: return "Worth Monitoring"
    elif t >= 40: return "Average Range"
    else:         return "Below Average"

def get_level_ar(t):
    if t >= 70:   return "ملحوظ بشكل واضح"
    elif t >= 65: return "مصدر قلق محتمل"
    elif t >= 60: return "يستحق المتابعة"
    elif t >= 40: return "ضمن المتوسط"
    else:         return "أقل من المتوسط"

def get_bar_color(t):
    if t >= 70:   return "#D32F2F"
    elif t >= 65: return "#F57C00"
    elif t >= 60: return "#FBC02D"
    elif t >= 40: return "#388E3C"
    else:         return "#1976D2"

def raw_to_t(raw, key):
    mean, sd = NORMS[key]
    if sd == 0: return 50
    return max(20, min(90, round(50 + 10 * (raw - mean) / sd)))

def compute_scores(responses):
    results = {}
    for key, info in SUBSCALES.items():
        raw = sum(responses.get(i, 0) for i in info["items"])
        results[key] = {"raw": raw, "t": raw_to_t(raw, key), "max_raw": len(info["items"]) * 3}
    return results

# ══════════════════════════════════════════════════════════════
#  CHARTS
# ══════════════════════════════════════════════════════════════
def make_bar_chart(scores, lang):
    labels, t_vals, colors_ = [], [], []
    for key in "ABCDEFGHIJKLMN":
        info = SUBSCALES[key]
        t    = scores[key]["t"]
        # Always use English labels in the chart to avoid font/rendering issues
        labels.append(info["name_en"])
        t_vals.append(t); colors_.append(get_bar_color(t))

    fig, ax = plt.subplots(figsize=(12, 7))
    fig.patch.set_facecolor('#F7F3EE'); ax.set_facecolor('#F7F3EE')
    y_pos = np.arange(len(labels))
    bars  = ax.barh(y_pos, t_vals, color=colors_, edgecolor='white', linewidth=0.8, height=0.65)
    for xv, lbl, col in [(40,'T=40','#388E3C'),(60,'T=60','#FBC02D'),
                          (65,'T=65','#F57C00'),(70,'T=70','#D32F2F')]:
        ax.axvline(x=xv, color=col, linestyle='--', linewidth=1.2, alpha=0.7, label=lbl)
    for bar_, val in zip(bars, t_vals):
        ax.text(bar_.get_width()+0.5, bar_.get_y()+bar_.get_height()/2,
                str(val), va='center', ha='left', fontsize=9, fontweight='bold', color='#1C1917')
    ax.set_yticks(y_pos); ax.set_yticklabels(labels, fontsize=9.5, fontfamily='DejaVu Sans')
    ax.set_xlim(20, 95)
    ax.set_xlabel('T-Score', fontsize=11, fontweight='bold', color='#1C1917')
    title = "Conners' CPRS-R:L — T-Score Profile"
    ax.set_title(title, fontsize=13, fontweight='bold', color='#1C1917', pad=14)
    ax.legend(loc='lower right', fontsize=8.5, framealpha=0.7)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='x', linestyle=':', alpha=0.5)
    ax.axvspan(70, 95, alpha=0.06, color='#D32F2F'); ax.axvspan(65, 70, alpha=0.05, color='#F57C00')
    plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig); buf.seek(0)
    return buf.read()

def make_pie_chart(responses):
    counts = [0,0,0,0]
    for v in responses.values(): counts[v] += 1
    labels  = ['0 — Not at all','1 — Just a little','2 — Pretty much','3 — Very much']
    colors_ = ['#388E3C','#FBC02D','#F57C00','#D32F2F']
    fig, ax = plt.subplots(figsize=(6, 4.5))
    fig.patch.set_facecolor('#F7F3EE')
    wedges, texts, autotexts = ax.pie(counts, labels=labels, colors=colors_,
        autopct='%1.0f%%', startangle=90, wedgeprops={'edgecolor':'white','linewidth':1.5})
    for at in autotexts: at.set_fontsize(9); at.set_fontweight('bold')
    ax.set_title('Response Distribution', fontsize=11, fontweight='bold', color='#1C1917')
    plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig); buf.seek(0)
    return buf.read()

# ══════════════════════════════════════════════════════════════
#  GROQ REPORTS
# ══════════════════════════════════════════════════════════════
def build_score_block_en(scores):
    return "\n".join(
        f"  {k}. {SUBSCALES[k]['name_en']}: Raw={scores[k]['raw']}/{scores[k]['max_raw']}, T={scores[k]['t']} — {get_level_en(scores[k]['t'])}"
        for k in "ABCDEFGHIJKLMN"
    )

def generate_report_en(child_name, age, gender, rater, scores):
    elevated = [k for k in "ABCDEFGHIJKLMN" if scores[k]["t"] >= 65]
    adhd_index_t = scores["H"]["t"]
    inatt_t      = scores["L"]["t"]
    hyper_t      = scores["M"]["t"]
    total_adhd_t = scores["N"]["t"]
    cgi_total_t  = scores["K"]["t"]
    prompt = f"""You are a licensed child psychologist writing a professional CPRS-R:L assessment report.

CHILD: {child_name} | AGE: {age} | GENDER: {gender} | RATER: {rater}
ASSESSMENT: Conners' Parent Rating Scale – Revised: Long Version (CPRS-R:L)
DATE: {date.today().strftime('%B %d, %Y')}

SUBSCALE T-SCORES (T≥65 = clinically significant; T≥70 = markedly atypical):
{build_score_block_en(scores)}

ELEVATED SCALES (T≥65): {', '.join(elevated) if elevated else 'None'}

KEY ADHD INDICATORS:
- ADHD Index (H): T={adhd_index_t}
- DSM-IV Inattentive (L): T={inatt_t}
- DSM-IV Hyperactive-Impulsive (M): T={hyper_t}
- DSM-IV Total (N): T={total_adhd_t}
- CGI Total (K): T={cgi_total_t}

RULES:
- Do NOT diagnose. State findings as hypotheses requiring clinical judgment.
- Use formal clinical language.
- Be specific to the T-scores above.
- No markdown symbols (**, ##, ---).
- Section titles: ALL CAPS numbered. Example: 1. REFERRAL & ASSESSMENT OVERVIEW
- The primary focus of the report is ADHD and related symptoms (inattention, hyperactivity, impulsivity).
  Prioritize and expand ADHD subscale discussion. Identify whether the profile suggests primarily
  inattentive, primarily hyperactive/impulsive, or combined presentation.

REPORT STRUCTURE:

CONNERS' PARENT RATING SCALE — CLINICAL REPORT
Child | {child_name}
Age | {age}  |  Gender | {gender}
Rater | {rater}
Date | {date.today().strftime('%B %d, %Y')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

CLINICAL SUMMARY
3–5 sentences covering the ADHD index, most elevated scales, and clinical significance.

1. REFERRAL & ASSESSMENT OVERVIEW
Instrument description, purpose, administration context, rating period (past month).

2. ADHD SUBSCALE ANALYSIS
Detailed analysis of scales H, L, M, N. Identify the ADHD presentation pattern (inattentive /
hyperactive-impulsive / combined). Interpret T-scores and behavioral correlates in depth.

3. OTHER CONTENT SUBSCALES PROFILE
For each non-ADHD scale T≥60: dedicated paragraph with T-score, behavioral correlates, clinical significance.
For T<60: one brief line noting within-normal-limits finding.

4. CLINICAL GLOBAL INDEX (I, J, K)
Interpret the CGI scales. Discuss overall severity of behavioral concerns.

5. STRENGTHS & PROTECTIVE FACTORS
Identify subscales in average or below-average range as relative strengths.

6. INTEGRATED CLINICAL IMPRESSIONS
Synthesize the ADHD profile. What overall pattern emerges? Primary areas of concern?

7. RECOMMENDATIONS
Evidence-based recommendations for ADHD intervention, monitoring, referral, or further assessment.

8. SUMMARY
One paragraph for clinical records:
"According to the CPRS-R:L completed by {rater}, {child_name} (age {age}) presents with..."
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=3500
    )
    return r.choices[0].message.content.strip()

def generate_report_ar(child_name, age, gender, rater, scores):
    elevated_ar = [SUBSCALES[k]["name_ar"] for k in "ABCDEFGHIJKLMN" if scores[k]["t"] >= 65]
    score_block = "\n".join(
        f"  {k}. {SUBSCALES[k]['name_ar']}: خام={scores[k]['raw']}/{scores[k]['max_raw']}, تائي={scores[k]['t']} — {get_level_ar(scores[k]['t'])}"
        for k in "ABCDEFGHIJKLMN"
    )
    # Determine likely ADHD pattern from subscale scores
    adhd_index_t   = scores["H"]["t"]
    inatt_t        = scores["L"]["t"]
    hyper_t        = scores["M"]["t"]
    total_adhd_t   = scores["N"]["t"]
    cgi_total_t    = scores["K"]["t"]

    prompt = f"""أنت طبيب نفسي للأطفال تكتب تقريراً سريرياً احترافياً لمقياس كونرز للوالدين (النسخة المراجعة الطويلة).

الطفل: {child_name} (يُكتب بالعربية: قدِّم تعريبًا صوتيًا مناسبًا للاسم) | السن: {age} | النوع: {gender} | المُقيِّم: {rater}
المقياس: مقياس كونرز للوالدين — نسخة مراجعة طويلة (CPRS-R:L)
التاريخ: {date.today().strftime('%Y/%m/%d')}

الدرجات التائية (T≥65 = ذو دلالة سريرية؛ T≥70 = ملحوظ بشكل واضح):
{score_block}

المقاييس المرتفعة (T≥65): {', '.join(elevated_ar) if elevated_ar else 'لا يوجد'}

مؤشرات ADHD الرئيسية:
- مؤشر ADHD (H): T={adhd_index_t}
- نقص الانتباه DSM-IV (L): T={inatt_t}
- فرط الحركة والاندفاعية DSM-IV (M): T={hyper_t}
- المجموع الكلي DSM-IV (N): T={total_adhd_t}
- المؤشر السريري العام-المجموع (K): T={cgi_total_t}

قواعد صارمة:
- لا تضع تشخيصاً. أشر إلى النتائج كفرضيات تحتاج إلى حكم سريري.
- استخدم لغة سريرية رسمية بالعربية الكاملة. لا إنجليزية إلا للاختصارات الطبية (CPRS-R:L, DSM-IV, ADHD, CGI, T-score).
- لا رموز markdown (**, ##, ---).
- عناوين الأقسام: أرقام + عنوان. مثال: ١. نظرة عامة على التقييم
- التركيز الأساسي للتقرير هو اضطراب ADHD وأعراضه (نقص الانتباه، فرط الحركة، الاندفاعية) والمقاييس ذات الصلة.
- عند ذكر اسم الطفل في التقرير، استخدم التعريب الصوتي العربي للاسم الإنجليزي المُدخل.

هيكل التقرير:

تقرير مقياس كونرز للوالدين — التقرير السريري
الطفل | {child_name}
السن | {age}  |  النوع | {gender}
المُقيِّم | {rater}
التاريخ | {date.today().strftime('%Y/%m/%d')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

ملخص سريري
٣–٥ جمل: تلخيص المقياس العام مع التركيز على مؤشرات ADHD، المقاييس الأكثر ارتفاعاً، والدلالة السريرية.

١. نظرة عامة على التقييم
وصف المقياس، الغرض منه، سياق التطبيق، فترة التقييم (الشهر الماضي).

٢. تحليل مؤشرات ADHD
ناقش بالتفصيل: مؤشر ADHD (H)، مقياس نقص الانتباه DSM-IV (L)، مقياس فرط الحركة والاندفاعية DSM-IV (M)، المجموع الكلي (N).
حدد النمط المحتمل: هل الأعراض أكثر في نقص الانتباه، أم فرط الحركة/الاندفاعية، أم كليهما؟

٣. تحليل المقاييس الفرعية الأخرى
لكل مقياس T≥60: فقرة مخصصة بالدرجة التائية والمؤشرات السلوكية والدلالة السريرية.
للمقاييس T<60: سطر واحد موجز يشير إلى أنها ضمن الحدود الطبيعية.

٤. المؤشر السريري العام (CGI)
فسّر مقاييس CGI (I: الاندفاعية وعدم الهدوء، J: التقلب الانفعالي، K: المجموع العام).

٥. نقاط القوة والعوامل الوقائية
حدد المقاييس ضمن المتوسط أو دونه كنقاط قوة نسبية.

٦. الانطباعات السريرية المتكاملة
ما النمط الكلي الذي يبرز؟ ما المجالات الأساسية للقلق من منظور ADHD؟

٧. التوصيات
توصيات مبنية على الأدلة للتدخل، المتابعة، الإحالة، أو التقييم الإضافي — مع التركيز على ADHD.

٨. الملخص — فقرة واحدة مناسبة للسجلات السريرية.
"وفقاً لمقياس CPRS-R:L الذي أكمله {rater}، يُظهر الطفل (التعريب الصوتي للاسم {child_name}) بعمر {age}..."
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=3500
    )
    return r.choices[0].message.content.strip()

# ══════════════════════════════════════════════════════════════
#  PDF REPORT BUILDER  (English only)
# ══════════════════════════════════════════════════════════════
# Color palette
PDF_WARM    = colors.HexColor('#8B7355')
PDF_DARK    = colors.HexColor('#1C1917')
PDF_CREAM   = colors.HexColor('#F7F3EE')
PDF_BORDER  = colors.HexColor('#DDD5C8')
PDF_ACCENT  = colors.HexColor('#C4956A')
PDF_RED     = colors.HexColor('#C62828')
PDF_ORANGE  = colors.HexColor('#F57C00')
PDF_YELLOW  = colors.HexColor('#FBC02D')
PDF_GREEN   = colors.HexColor('#388E3C')
PDF_HEADER  = colors.HexColor('#2D2926')

def _t_band_color(t):
    if t >= 70: return PDF_RED
    if t >= 65: return PDF_ORANGE
    if t >= 60: return PDF_YELLOW
    return PDF_GREEN

def _t_band_label(t):
    if t >= 70: return "Markedly Atypical"
    if t >= 65: return "Likely Concern"
    if t >= 60: return "Worth Monitoring"
    if t >= 40: return "Average Range"
    return "Below Average"

def _make_pdf_styles():
    base = getSampleStyleSheet()
    styles = {}
    styles['title'] = ParagraphStyle('title', fontName='Helvetica-Bold',
        fontSize=16, textColor=PDF_DARK, spaceAfter=4, alignment=TA_CENTER)
    styles['subtitle'] = ParagraphStyle('subtitle', fontName='Helvetica',
        fontSize=9, textColor=PDF_WARM, spaceAfter=2, alignment=TA_CENTER)
    styles['section'] = ParagraphStyle('section', fontName='Helvetica-Bold',
        fontSize=11, textColor=PDF_WARM, spaceBefore=14, spaceAfter=4)
    styles['body'] = ParagraphStyle('body', fontName='Helvetica',
        fontSize=9.5, textColor=PDF_DARK, leading=14, spaceAfter=5)
    styles['small'] = ParagraphStyle('small', fontName='Helvetica',
        fontSize=8, textColor=PDF_WARM, leading=11)
    styles['bold_body'] = ParagraphStyle('bold_body', fontName='Helvetica-Bold',
        fontSize=9.5, textColor=PDF_DARK, leading=14, spaceAfter=3)
    styles['summary_box'] = ParagraphStyle('summary_box', fontName='Helvetica',
        fontSize=9.5, textColor=PDF_DARK, leading=14, spaceAfter=0,
        leftIndent=6, rightIndent=6)
    return styles

def build_pdf_report_en(report_text, scores, bar_bytes, pie_bytes,
                         child_name, age, gender, rater, responses_dict):
    """Build English PDF report. Returns BytesIO."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm)
    S = _make_pdf_styles()
    W = A4[0] - 4*cm   # usable width
    story = []

    # ── Logo ──
    if os.path.exists(LOGO_FILE):
        try:
            logo = RLImage(LOGO_FILE, width=5*cm, height=2.2*cm)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 4))
        except: pass

    # ── Title block ──
    story.append(Paragraph("Conners' Parent Rating Scale — Clinical Report", S['title']))
    story.append(Paragraph("CPRS-R:L — Revised: Long Version", S['subtitle']))
    story.append(HRFlowable(width=W, thickness=1, color=PDF_BORDER, spaceAfter=10))

    # ── Demographics table ──
    demo_data = [
        [Paragraph('<b>Child</b>', S['small']),  Paragraph(child_name or '—', S['body']),
         Paragraph('<b>Age</b>', S['small']),    Paragraph(str(age) or '—', S['body'])],
        [Paragraph('<b>Gender</b>', S['small']), Paragraph(gender or '—', S['body']),
         Paragraph('<b>Rater</b>', S['small']),  Paragraph(rater or '—', S['body'])],
        [Paragraph('<b>Date</b>', S['small']),   Paragraph(date.today().strftime('%B %d, %Y'), S['body']),
         Paragraph('<b>Assessment</b>', S['small']), Paragraph('CPRS-R:L (80 items, 0–3)', S['body'])],
    ]
    demo_tbl = Table(demo_data, colWidths=[2.2*cm, 6.5*cm, 2.2*cm, 6.5*cm])
    demo_tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), PDF_CREAM),
        ('BOX', (0,0), (-1,-1), 0.5, PDF_BORDER),
        ('INNERGRID', (0,0), (-1,-1), 0.3, PDF_BORDER),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ('LEFTPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(demo_tbl)
    story.append(Spacer(1, 10))

    # ── Subscale score summary table ──
    story.append(Paragraph("SUBSCALE SCORE SUMMARY", S['section']))
    score_header = [
        Paragraph('<b>Scale</b>', S['small']),
        Paragraph('<b>Raw</b>', S['small']),
        Paragraph('<b>T-Score</b>', S['small']),
        Paragraph('<b>Classification</b>', S['small']),
        Paragraph('<b>Band</b>', S['small']),
    ]
    score_rows = [score_header]
    for key in "ABCDEFGHIJKLMN":
        s = scores[key]
        t = s['t']
        band_col = _t_band_color(t)
        score_rows.append([
            Paragraph(f"{key}. {SUBSCALES[key]['name_en']}", S['body']),
            Paragraph(f"{s['raw']}/{s['max_raw']}", S['body']),
            Paragraph(f"<b>{t}</b>", S['body']),
            Paragraph(_t_band_label(t), S['body']),
            Paragraph('', S['body']),  # colored cell via style
        ])
    score_tbl = Table(score_rows, colWidths=[6.5*cm, 1.8*cm, 1.8*cm, 4*cm, 3.4*cm])
    ts_cmds = [
        ('BACKGROUND', (0,0), (-1,0), PDF_HEADER),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8.5),
        ('BOX', (0,0), (-1,-1), 0.5, PDF_BORDER),
        ('INNERGRID', (0,0), (-1,-1), 0.3, PDF_BORDER),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, PDF_CREAM]),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
    ]
    for row_i, key in enumerate("ABCDEFGHIJKLMN", start=1):
        t = scores[key]['t']
        ts_cmds.append(('BACKGROUND', (4, row_i), (4, row_i), _t_band_color(t)))
        ts_cmds.append(('TEXTCOLOR',  (4, row_i), (4, row_i), colors.black))
        ts_cmds.append(('FONTNAME',   (4, row_i), (4, row_i), 'Helvetica-Bold'))
    score_tbl.setStyle(TableStyle(ts_cmds))
    story.append(score_tbl)
    story.append(Spacer(1, 6))

    # ── Colour legend ──
    legend_data = [
        [Paragraph('<b>Colour</b>', S['small']),
         Paragraph('<b>T-Score Range</b>', S['small']),
         Paragraph('<b>Classification</b>', S['small']),
         Paragraph('<b>Clinical Meaning</b>', S['small'])],
        [Paragraph('', S['small']),
         Paragraph('T ≥ 70', S['body']),
         Paragraph('<b>Markedly Atypical</b>', S['body']),
         Paragraph('Strong clinical concern. Thorough assessment and intervention strongly recommended.', S['body'])],
        [Paragraph('', S['small']),
         Paragraph('65 ≤ T < 70', S['body']),
         Paragraph('<b>Likely Concern</b>', S['body']),
         Paragraph('Clinically significant. Warrants follow-up and further clinical evaluation.', S['body'])],
        [Paragraph('', S['small']),
         Paragraph('60 ≤ T < 65', S['body']),
         Paragraph('<b>Worth Monitoring</b>', S['body']),
         Paragraph('Mildly elevated. Monitor over time; intervention may be beneficial.', S['body'])],
        [Paragraph('', S['small']),
         Paragraph('40 ≤ T < 60', S['body']),
         Paragraph('<b>Average Range</b>', S['body']),
         Paragraph('Within normal limits. No clinical concern indicated.', S['body'])],
    ]
    legend_tbl = Table(legend_data, colWidths=[1.2*cm, 3*cm, 4*cm, W-8.2*cm])
    legend_cmds = [
        ('BACKGROUND', (0,0), (-1,0), PDF_HEADER),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0,0), (-1,-1), 8.5),
        ('BOX',        (0,0), (-1,-1), 0.5, PDF_BORDER),
        ('INNERGRID',  (0,0), (-1,-1), 0.3, PDF_BORDER),
        ('VALIGN',     (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING',(0,0),(-1,-1), 4),
        ('LEFTPADDING',(0,0), (-1,-1), 5),
        # colour swatch cells — col 0 of each data row
        ('BACKGROUND', (0,1), (0,1), PDF_RED),
        ('BACKGROUND', (0,2), (0,2), PDF_ORANGE),
        ('BACKGROUND', (0,3), (0,3), PDF_YELLOW),
        ('BACKGROUND', (0,4), (0,4), PDF_GREEN),
        # text on coloured swatches → black bold
        ('TEXTCOLOR',  (0,1), (0,-1), colors.black),
        ('FONTNAME',   (0,1), (0,-1), 'Helvetica-Bold'),
        # alternate row background for readability cols
        ('ROWBACKGROUNDS', (1,1), (-1,-1), [colors.white, PDF_CREAM, colors.white, PDF_CREAM]),
    ]
    legend_tbl.setStyle(TableStyle(legend_cmds))
    story.append(Paragraph("COLOUR LEGEND — T-SCORE CLASSIFICATION", S['section']))
    story.append(legend_tbl)
    story.append(Spacer(1, 8))

    # ── Bar chart ──
    story.append(Paragraph("T-SCORE PROFILE CHART", S['section']))
    bar_img = RLImage(io.BytesIO(bar_bytes), width=W, height=W*0.52)
    bar_img.hAlign = 'CENTER'
    story.append(bar_img)
    story.append(Spacer(1, 6))

    # ── Pie + ADHD highlight side by side ──
    story.append(Paragraph("RESPONSE DISTRIBUTION & ADHD INDICATORS", S['section']))
    pie_img = RLImage(io.BytesIO(pie_bytes), width=7.5*cm, height=5.5*cm)

    adhd_keys = [("H","ADHD Index"),("L","DSM-IV Inattentive"),
                 ("M","DSM-IV Hyperactive-Impulsive"),("N","DSM-IV Total")]
    adhd_rows = [[Paragraph('<b>Scale</b>',S['small']),
                  Paragraph('<b>T</b>',S['small']),
                  Paragraph('<b>Classification</b>',S['small'])]]
    for k, lbl in adhd_keys:
        t = scores[k]['t']
        adhd_rows.append([Paragraph(lbl, S['body']),
                          Paragraph(f"<b>{t}</b>", S['body']),
                          Paragraph(_t_band_label(t), S['body'])])
    adhd_tbl = Table(adhd_rows, colWidths=[4.5*cm, 1.2*cm, 3.8*cm])
    adhd_ts = [
        ('BACKGROUND', (0,0), (-1,0), PDF_HEADER),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8.5),
        ('BOX', (0,0), (-1,-1), 0.5, PDF_BORDER),
        ('INNERGRID', (0,0), (-1,-1), 0.3, PDF_BORDER),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, PDF_CREAM]),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 4),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ('LEFTPADDING', (0,0), (-1,-1), 5),
    ]
    for row_i, (k,_) in enumerate(adhd_keys, start=1):
        t = scores[k]['t']
        adhd_ts.append(('BACKGROUND', (1, row_i), (1, row_i), _t_band_color(t)))
        adhd_ts.append(('TEXTCOLOR',  (1, row_i), (1, row_i), colors.black))
        adhd_ts.append(('FONTNAME',   (1, row_i), (1, row_i), 'Helvetica-Bold'))
    adhd_tbl.setStyle(TableStyle(adhd_ts))

    side_tbl = Table([[pie_img, adhd_tbl]], colWidths=[8*cm, W-8*cm])
    side_tbl.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                                   ('LEFTPADDING',(1,0),(1,0),12)]))
    story.append(side_tbl)
    story.append(Spacer(1, 6))

    # ── Clinical narrative ──
    story.append(HRFlowable(width=W, thickness=0.5, color=PDF_BORDER))
    story.append(Paragraph("CLINICAL NARRATIVE REPORT", S['section']))
    sec_en_pat = re.compile(r'^\d+\.\s+[A-Z][A-Z\s&/()\-:]+$')
    header_words = {"CONNERS' PARENT RATING SCALE — CLINICAL REPORT","CLINICAL SUMMARY"}

    for line in report_text.split('\n'):
        ls = line.strip()
        if not ls:
            story.append(Spacer(1, 4)); continue
        if ls.startswith('━') or ls.startswith('═'):
            story.append(HRFlowable(width=W, thickness=0.4, color=PDF_BORDER, spaceAfter=4)); continue
        is_section = (sec_en_pat.match(ls) or ls in header_words or
                      ls.upper() in header_words or ls == "CLINICAL SUMMARY")
        if is_section:
            story.append(Paragraph(ls, S['section'])); continue
        # inline table rows (pipe-separated)
        if '|' in ls:
            parts = [p.strip() for p in ls.split('|') if p.strip()]
            if len(parts) >= 2:
                skip = [("field","value"),("subscale","raw")]
                if (parts[0].lower(), parts[1].lower()) not in skip:
                    row_data = [[Paragraph(p, S['body']) for p in parts]]
                    col_w = W / len(parts)
                    mini_tbl = Table(row_data, colWidths=[col_w]*len(parts))
                    mini_tbl.setStyle(TableStyle([
                        ('BOX',(0,0),(-1,-1),0.3,PDF_BORDER),
                        ('INNERGRID',(0,0),(-1,-1),0.3,PDF_BORDER),
                        ('BACKGROUND',(0,0),(-1,-1),PDF_CREAM),
                        ('TOPPADDING',(0,0),(-1,-1),3),
                        ('BOTTOMPADDING',(0,0),(-1,-1),3),
                        ('LEFTPADDING',(0,0),(-1,-1),5),
                    ]))
                    story.append(mini_tbl); continue
        story.append(Paragraph(ls, S['body']))

    story.append(Spacer(1, 10))
    story.append(HRFlowable(width=W, thickness=0.5, color=PDF_BORDER))

    # ── Item responses table (~2 pages) ──
    story.append(PageBreak())
    story.append(Paragraph("ITEM RESPONSES — FULL RATING TABLE", S['section']))
    story.append(Paragraph(
        "The following table shows each item number, the item text, and the rating assigned by the rater. "
        "Rating key: 0 = Not at all · 1 = Just a little · 2 = Pretty much · 3 = Very much",
        S['small']))
    story.append(Spacer(1, 6))

    rating_labels = {0:"0 — Not at all", 1:"1 — Just a little",
                     2:"2 — Pretty much",  3:"3 — Very much"}
    rating_colors = {0:PDF_GREEN, 1:PDF_YELLOW, 2:PDF_ORANGE, 3:PDF_RED}

    item_header = [
        Paragraph('<b>#</b>', S['small']),
        Paragraph('<b>Item</b>', S['small']),
        Paragraph('<b>Rating</b>', S['small']),
    ]
    item_rows = [item_header]
    item_ts_cmds = [
        ('BACKGROUND', (0,0), (-1,0), PDF_HEADER),
        ('TEXTCOLOR', (0,0), (-1,0), colors.white),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 8),
        ('BOX', (0,0), (-1,-1), 0.5, PDF_BORDER),
        ('INNERGRID', (0,0), (-1,-1), 0.3, PDF_BORDER),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
    ]
    for i, item_text in enumerate(ITEMS_EN):
        item_num = i + 1
        val = responses_dict.get(item_num, 0)
        bg = colors.white if i % 2 == 0 else PDF_CREAM
        item_ts_cmds.append(('BACKGROUND', (0, item_num), (1, item_num), bg))
        item_ts_cmds.append(('BACKGROUND', (2, item_num), (2, item_num), rating_colors[val]))
        item_ts_cmds.append(('TEXTCOLOR',  (2, item_num), (2, item_num), colors.black))
        item_ts_cmds.append(('FONTNAME',   (2, item_num), (2, item_num), 'Helvetica-Bold'))
        item_rows.append([
            Paragraph(str(item_num), S['small']),
            Paragraph(item_text, S['small']),
            Paragraph(rating_labels[val], S['small']),
        ])
    item_tbl = Table(item_rows, colWidths=[1*cm, 12.5*cm, 4*cm])
    item_tbl.setStyle(TableStyle(item_ts_cmds))
    story.append(item_tbl)

    # ── Confidentiality footer ──
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width=W, thickness=0.5, color=PDF_BORDER))
    story.append(Spacer(1, 4))
    story.append(Paragraph(
        "This report is strictly confidential. Scores are based on parent/caregiver rating and should be "
        "interpreted in conjunction with clinical judgment and other assessment data. "
        "CPRS-R:L T-scores ≥65 are considered clinically significant.",
        S['small']))

    doc.build(story)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════
#  WORD DOC BUILDER  (Arabic only)
# ══════════════════════════════════════════════════════════════
def build_word_report(report_text, scores, bar_bytes, pie_bytes,
                      child_name, age, gender, rater, lang, responses_dict=None):
    is_rtl = (lang == "ar")
    doc = Document()
    for sec_ in doc.sections:
        sec_.top_margin=Cm(2.0); sec_.bottom_margin=Cm(2.0)
        sec_.left_margin=Cm(2.2); sec_.right_margin=Cm(2.2)

    # Page border
    for sec_ in doc.sections:
        sp = sec_._sectPr; pb = OxmlElement('w:pgBorders')
        pb.set(qn('w:offsetFrom'),'page')
        for side in ('top','left','bottom','right'):
            b = OxmlElement(f'w:{side}'); b.set(qn('w:val'),'single')
            b.set(qn('w:sz'),'10'); b.set(qn('w:space'),'24')
            b.set(qn('w:color'),'8B7355'); pb.append(b)
        sp.append(pb)

    # Footer
    for sec_ in doc.sections:
        ft = sec_.footer
        fp = ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
        fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_ = fp.add_run(); r_.font.size=Pt(9); r_.font.color.rgb=WARM_RGB
        for tag, text in [('begin',None),(None,' PAGE '),('end',None)]:
            if tag:
                el=OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),tag); r_._r.append(el)
            else:
                it=OxmlElement('w:instrText'); it.text=text; r_._r.append(it)

    def set_rtl(p):
        if is_rtl:
            pPr=p._p.get_or_add_pPr()
            pPr.append(OxmlElement("w:bidi"))
            jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)

    def add_para(text, bold=False, size=11, color=None, space_before=0, space_after=4,
                 alignment=None, keep_next=False):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(space_before)
        p.paragraph_format.space_after=Pt(space_after)
        if keep_next: p.paragraph_format.keep_with_next=True
        set_rtl(p)
        if alignment: p.alignment=alignment
        r_=p.add_run(text); r_.font.size=Pt(size); r_.font.name="Times New Roman"; r_.font.bold=bold
        if color: r_.font.color.rgb=color
        return p

    def add_section_title(text):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(12); p.paragraph_format.space_after=Pt(3)
        p.paragraph_format.keep_with_next=True
        set_rtl(p)
        r_=p.add_run(text.strip()); r_.font.size=Pt(13); r_.font.name="Times New Roman"
        r_.font.bold=True; r_.font.color.rgb=WARM_RGB
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
        bot.set(qn('w:space'),'2'); bot.set(qn('w:color'),'C4956A')
        pBdr.append(bot); pPr.append(pBdr)

    def make_table():
        t=doc.add_table(rows=0, cols=2); t.style='Table Grid'
        try:
            tPr=t._tbl.tblPr
            if is_rtl:
                bv=OxmlElement('w:bidiVisual'); tPr.append(bv)
            tW=OxmlElement('w:tblW'); tW.set(qn('w:w'),'9026'); tW.set(qn('w:type'),'dxa'); tPr.append(tW)
            tg=OxmlElement('w:tblGrid')
            for w in [3000,6026]:
                gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); tg.append(gc)
            t._tbl.insert(0,tg)
        except: pass
        return t

    def add_row(table, field, value, header=False):
        row=table.add_row()
        trPr=row._tr.get_or_add_trPr()
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trPr.append(cs)
        if is_rtl:
            bd=OxmlElement('w:bidi'); trPr.append(bd)
        for idx,(cell,txt,bold_) in enumerate([(row.cells[0],field,True),(row.cells[1],value,header)]):
            cell.text=""
            p=cell.paragraphs[0]; set_rtl(p)
            lines_=str(txt).split('\n') if '\n' in str(txt) else [str(txt)]
            for li,line_ in enumerate(lines_):
                vp=p if li==0 else cell.add_paragraph()
                vr=vp.add_run(line_.strip()); vr.font.size=Pt(10); vr.font.name="Times New Roman"; vr.font.bold=bold_
                if header: vr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            tc=cell._tc; tcP=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            if header:
                shd.set(qn('w:fill'),'2D2926' if idx==0 else '8B7355')
            elif idx==0:
                shd.set(qn('w:fill'),'F7F3EE')
            else:
                shd.set(qn('w:fill'),'FFFFFF')
            tcP.append(shd)
            mg=OxmlElement('w:tcMar')
            for side in ['top','bottom','left','right']:
                m=OxmlElement(f'w:{side}'); m.set(qn('w:w'),'60'); m.set(qn('w:type'),'dxa'); mg.append(m)
            tcP.append(mg)

    # ── Header ──
    p_hdr=doc.add_paragraph(); p_hdr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_before=Pt(0); p_hdr.paragraph_format.space_after=Pt(6)
    if os.path.exists(LOGO_FILE):
        p_hdr.add_run().add_picture(LOGO_FILE, width=Inches(2.8))

    title_text = ("Conners' Parent Rating Scale — Clinical Report" if lang=="en"
                  else "تقرير مقياس كونرز للوالدين — التقرير السريري")
    r_t=p_hdr.add_run(f"\n{title_text}")
    r_t.font.name="Times New Roman"; r_t.font.size=Pt(17)
    r_t.font.bold=True; r_t.font.color.rgb=DARK_RGB

    sub_text=("CPRS-R:L — Conners' Parent Rating Scale — Revised: Long Version" if lang=="en"
              else "CPRS-R:L — مقياس كونرز للوالدين — نسخة مراجعة طويلة")
    add_para(sub_text, size=9.5, color=WARM_RGB,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    p_sep=doc.add_paragraph(); p_sep.paragraph_format.space_before=Pt(2); p_sep.paragraph_format.space_after=Pt(8)
    pPr=p_sep._p.get_or_add_pPr(); pBdr2=OxmlElement('w:pBdr')
    bot2=OxmlElement('w:bottom'); bot2.set(qn('w:val'),'single'); bot2.set(qn('w:sz'),'8')
    bot2.set(qn('w:space'),'2'); bot2.set(qn('w:color'),'DDD5C8')
    pBdr2.append(bot2); pPr.append(pBdr2)

    # Client info table
    info_tbl=make_table()
    if lang=="en":
        add_row(info_tbl,"Field","Value",header=True)
        add_row(info_tbl,"Child",child_name); add_row(info_tbl,"Age",age)
        add_row(info_tbl,"Gender",gender);   add_row(info_tbl,"Rater",rater)
        add_row(info_tbl,"Date",date.today().strftime('%B %d, %Y'))
        add_row(info_tbl,"Assessment","CPRS-R:L (80 items, scale 0–3)")
    else:
        add_row(info_tbl,"الحقل","البيانات",header=True)
        add_row(info_tbl,"الطفل",child_name); add_row(info_tbl,"السن",age)
        add_row(info_tbl,"النوع",gender);     add_row(info_tbl,"المُقيِّم",rater)
        add_row(info_tbl,"التاريخ",date.today().strftime('%Y/%m/%d'))
        add_row(info_tbl,"المقياس","CPRS-R:L (80 بنداً، مقياس 0–3)")
    doc.add_paragraph().paragraph_format.space_after=Pt(4)

    # ── Color legend ──
    legend_title = "COLOUR LEGEND — T-SCORE CLASSIFICATION" if lang=="en" else "دليل الألوان — تصنيف الدرجات التائية"
    add_section_title(legend_title)
    legend_tbl2 = doc.add_table(rows=0, cols=3); legend_tbl2.style = 'Table Grid'
    legend_entries = [
        ("T ≥ 70",    "Markedly Atypical"  if lang=="en" else "ملحوظ بشكل واضح",  "Strong clinical concern. Thorough assessment and intervention strongly recommended."  if lang=="en" else "مصدر قلق سريري شديد. يُوصى بتقييم شامل وتدخل علاجي.", "FFCDD2"),
        ("65–69",     "Likely Concern"     if lang=="en" else "مصدر قلق محتمل",   "Clinically significant. Warrants follow-up and further evaluation."                   if lang=="en" else "ذو دلالة سريرية. يستدعي متابعة وتقييماً إضافياً.",   "FFE0B2"),
        ("60–64",     "Worth Monitoring"   if lang=="en" else "يستحق المتابعة",   "Mildly elevated. Monitor over time; intervention may be beneficial."                  if lang=="en" else "مرتفع قليلاً. يُتابع مع الوقت؛ قد يستفيد من التدخل.",  "FFF9C4"),
        ("40–59",     "Average Range"      if lang=="en" else "ضمن المتوسط",      "Within normal limits. No clinical concern indicated."                                  if lang=="en" else "ضمن الحدود الطبيعية. لا مخاوف سريرية.",              "D4EDDA"),
    ]
    # header
    hrow2 = legend_tbl2.add_row()
    hdrs2 = (["T-Score","Classification","Meaning"] if lang=="en"
             else ["الدرجة التائية","التصنيف","المعنى السريري"])
    for ci2, htxt2 in enumerate(hdrs2):
        cell2 = hrow2.cells[ci2]; cell2.text=""
        p2_ = cell2.paragraphs[0]; set_rtl(p2_)
        r2_ = p2_.add_run(htxt2); r2_.font.bold=True; r2_.font.size=Pt(9)
        r2_.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r2_.font.name="Times New Roman"
        tc2_ = cell2._tc; tcP2_ = tc2_.get_or_add_tcPr()
        shd2_ = OxmlElement('w:shd'); shd2_.set(qn('w:val'),'clear')
        shd2_.set(qn('w:color'),'auto'); shd2_.set(qn('w:fill'),'2D2926'); tcP2_.append(shd2_)
    for t_range, classif, meaning, fill_hex in legend_entries:
        lrow = legend_tbl2.add_row()
        for ci2, (cell_txt2, cell_fill) in enumerate([
            (t_range, fill_hex), (classif, fill_hex), (meaning, "FFFFFF")
        ]):
            cell2 = lrow.cells[ci2]; cell2.text=""
            p2_ = cell2.paragraphs[0]; set_rtl(p2_)
            r2_ = p2_.add_run(cell_txt2)
            r2_.font.bold = (ci2 < 2); r2_.font.size=Pt(9); r2_.font.name="Times New Roman"
            r2_.font.color.rgb = RGBColor(0,0,0)  # always black for readability
            tc2_ = cell2._tc; tcP2_ = tc2_.get_or_add_tcPr()
            shd2_ = OxmlElement('w:shd'); shd2_.set(qn('w:val'),'clear')
            shd2_.set(qn('w:color'),'auto'); shd2_.set(qn('w:fill'), cell_fill); tcP2_.append(shd2_)
            mg2_ = OxmlElement('w:tcMar')
            for side2 in ['top','bottom','left','right']:
                m2_ = OxmlElement(f'w:{side2}'); m2_.set(qn('w:w'),'60'); m2_.set(qn('w:type'),'dxa'); mg2_.append(m2_)
            tcP2_.append(mg2_)
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # Score summary
    sec_title="SUBSCALE SCORE SUMMARY" if lang=="en" else "ملخص درجات المقاييس الفرعية"
    add_section_title(sec_title)
    score_tbl=make_table()
    if lang=="en":
        add_row(score_tbl,"Subscale","Raw | T-Score | Classification",header=True)
        for key in "ABCDEFGHIJKLMN":
            s=scores[key]
            add_row(score_tbl,f"{key}. {SUBSCALES[key]['name_en']}",
                    f"{s['raw']}/{s['max_raw']}  |  T={s['t']}  |  {get_level_en(s['t'])}")
    else:
        add_row(score_tbl,"المقياس الفرعي","الخام | التائي | التصنيف",header=True)
        for key in "ABCDEFGHIJKLMN":
            s=scores[key]
            add_row(score_tbl,f"{key}. {SUBSCALES[key]['name_ar']}",
                    f"{s['raw']}/{s['max_raw']}  |  T={s['t']}  |  {get_level_ar(s['t'])}")
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # Bar chart
    add_section_title("T-SCORE PROFILE CHART" if lang=="en" else "مخطط الدرجات التائية")
    p_chart=doc.add_paragraph(); p_chart.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p_chart.paragraph_format.space_after=Pt(6)
    p_chart.add_run().add_picture(io.BytesIO(bar_bytes), width=Inches(6.2))

    # Pie chart
    add_section_title("RESPONSE DISTRIBUTION" if lang=="en" else "توزيع الاستجابات")
    p_pie=doc.add_paragraph(); p_pie.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p_pie.paragraph_format.space_after=Pt(6)
    p_pie.add_run().add_picture(io.BytesIO(pie_bytes), width=Inches(4.0))

    # Narrative
    add_section_title("CLINICAL NARRATIVE REPORT" if lang=="en" else "التقرير السريري التفصيلي")
    sec_en_pat=re.compile(r'^\d+\.\s+[A-Z][A-Z\s&/\(\):]+$')
    sec_ar_pat=re.compile(r'^[١٢٣٤٥٦٧٨٩\d]+[\.،:]\s+[\u0600-\u06FF]')
    header_words={"CONNERS' PARENT RATING SCALE — CLINICAL REPORT","CLINICAL SUMMARY",
                  "تقرير مقياس كونرز للوالدين","ملخص سريني","ملخص سريري"}
    in_table=False; current_table=None

    for line in report_text.split('\n'):
        ls=line.strip()
        if not ls:
            if in_table: in_table=False; current_table=None
            doc.add_paragraph().paragraph_format.space_after=Pt(2)
            continue
        is_section=(sec_en_pat.match(ls) or sec_ar_pat.match(ls) or
                    ls in header_words or ls.upper() in header_words)
        if is_section:
            in_table=False; current_table=None; add_section_title(ls); continue
        if ls.startswith('━') or ls.startswith('═'):
            in_table=False; current_table=None
            p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
            b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
            b.set(qn('w:space'),'1'); b.set(qn('w:color'),'DDD5C8')
            pBdr.append(b); pPr.append(pBdr); continue
        if '|' in ls:
            parts=[p.strip() for p in ls.split('|') if p.strip()]
            if not parts: continue
            if all(set(p)<=set('-: ') for p in parts): continue
            skip=[("field","value"),("subscale","raw"),("المقياس","الخام"),("الحقل","البيانات")]
            if len(parts)>=2 and (parts[0].strip('* ').lower(),parts[1].strip('* ').lower()) in skip: continue
            if not in_table or current_table is None:
                in_table=True; current_table=make_table()
                hdr=("Field","Details") if lang=="en" else ("الحقل","التفاصيل")
                add_row(current_table,hdr[0],hdr[1],header=True)
            field=parts[0].strip('* '); value=' | '.join(parts[1:])
            add_row(current_table,field,value); continue
        is_ar=any('\u0600'<=c<='\u06ff' for c in ls)
        if ls.endswith(':') and is_ar and len(ls)<60:
            in_table=False; current_table=None
            add_para(ls.rstrip(':'),bold=True,size=11,color=DARK_RGB,space_before=10,space_after=2,keep_next=True); continue
        in_table=False; current_table=None
        add_para(ls,size=10.5,space_before=0,space_after=3)

    doc.add_paragraph().paragraph_format.space_after=Pt(10)

    # ── Item responses table ──
    if responses_dict:
        doc.add_page_break()
        sec_title_items = "ITEM RESPONSES — FULL RATING TABLE" if lang=="en" else "جدول استجابات البنود الكاملة"
        add_section_title(sec_title_items)
        note_text = ("Rating key: 0 = Not at all · 1 = Just a little · 2 = Pretty much · 3 = Very much"
                     if lang=="en" else
                     "مفتاح التقييم: 0 = أبداً · 1 = أحياناً · 2 = إلى حد ما · 3 = كثيراً جداً")
        add_para(note_text, size=8.5, color=WARM_RGB, space_after=6)
        rating_labels_en = {0:"0 — Not at all",1:"1 — Just a little",2:"2 — Pretty much",3:"3 — Very much"}
        rating_labels_ar = {0:"0 — أبداً",1:"1 — أحياناً",2:"2 — إلى حد ما",3:"3 — كثيراً جداً"}
        rating_labels = rating_labels_en if lang=="en" else rating_labels_ar
        rating_fill   = {0:"D4EDDA",1:"FFF3CD",2:"FFE0B2",3:"FFCDD2"}
        ITEMS_TABEL   = ITEMS_EN if lang=="en" else ITEMS_AR
        item_tbl = doc.add_table(rows=0, cols=3)
        item_tbl.style = 'Table Grid'
        try:
            tPr2 = item_tbl._tbl.tblPr
            tW2  = OxmlElement('w:tblW'); tW2.set(qn('w:w'),'9026'); tW2.set(qn('w:type'),'dxa')
            tPr2.append(tW2)
        except: pass
        # header row
        hrow = item_tbl.add_row()
        for ci, htxt in enumerate(["#" if lang=="en" else "#",
                                    "Item" if lang=="en" else "البند",
                                    "Rating" if lang=="en" else "التقييم"]):
            cell = hrow.cells[ci]; cell.text=""
            p_ = cell.paragraphs[0]
            r_ = p_.add_run(htxt); r_.font.bold=True; r_.font.size=Pt(9)
            r_.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); r_.font.name="Times New Roman"
            tc_ = cell._tc; tcP_ = tc_.get_or_add_tcPr()
            shd_ = OxmlElement('w:shd'); shd_.set(qn('w:val'),'clear')
            shd_.set(qn('w:color'),'auto'); shd_.set(qn('w:fill'),'2D2926')
            tcP_.append(shd_)
        for i, item_text in enumerate(ITEMS_TABEL):
            item_num = i + 1
            val      = responses_dict.get(item_num, 0)
            irow     = item_tbl.add_row()
            fill_bg  = "F7F3EE" if i%2==0 else "FFFFFF"
            for ci, cell_txt in enumerate([str(item_num), item_text, rating_labels[val]]):
                cell = irow.cells[ci]; cell.text=""
                p_   = cell.paragraphs[0]
                if is_rtl and ci != 0:
                    pPr_ = p_._p.get_or_add_pPr()
                    pPr_.append(OxmlElement("w:bidi"))
                    jc_ = OxmlElement("w:jc"); jc_.set(qn("w:val"),"right"); pPr_.append(jc_)
                r_   = p_.add_run(cell_txt); r_.font.size=Pt(8.5); r_.font.name="Times New Roman"
                r_.font.bold = (ci == 2)  # bold on rating column
                r_.font.color.rgb = RGBColor(0, 0, 0)  # always black
                tc_  = cell._tc; tcP_ = tc_.get_or_add_tcPr()
                shd_ = OxmlElement('w:shd'); shd_.set(qn('w:val'),'clear'); shd_.set(qn('w:color'),'auto')
                if ci == 2:
                    shd_.set(qn('w:fill'), rating_fill[val])
                else:
                    shd_.set(qn('w:fill'), fill_bg)
                tcP_.append(shd_)
                mg_ = OxmlElement('w:tcMar')
                for side in ['top','bottom','left','right']:
                    m_ = OxmlElement(f'w:{side}'); m_.set(qn('w:w'),'50'); m_.set(qn('w:type'),'dxa'); mg_.append(m_)
                tcP_.append(mg_)
        doc.add_paragraph().paragraph_format.space_after=Pt(10)

    note=("This report is strictly confidential. Scores are based on parent/caregiver rating "
          "and should be interpreted in conjunction with clinical judgment and other assessment data. "
          "CPRS-R:L T-scores ≥65 are considered clinically significant.") if lang=="en" \
         else ("هذا التقرير سري للغاية. الدرجات مبنية على تقييم الوالدين وتُفسَّر بالتزامن مع "
               "الحكم السريري وبيانات التقييم الأخرى. الدرجات التائية ≥65 ذات دلالة سريرية.")
    add_para(note, size=8.5, color=WARM_RGB, space_before=6)

    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════
#  EMAIL
#  Arabic mode → EN PDF (word) + AR Word
#  English mode → EN Word only
# ══════════════════════════════════════════════════════════════
def send_email_ar(child_name, buf_pdf_en, buf_word_ar, fn_pdf_en, fn_word_ar, scores):
    """Arabic mode: attach English PDF + Arabic Word"""
    date_str=date.today().strftime('%B %d, %Y')
    elevated=[(k,scores[k]["t"]) for k in "ABCDEFGHIJKLMN" if scores[k]["t"]>=65]
    elev_html="".join(
        f"<tr><td style='padding:4px 0;color:#6B5B45;'>{SUBSCALES[k]['name_en']}</td>"
        f"<td><strong style='color:#D9534F;'>T={t}</strong></td></tr>"
        for k,t in elevated
    ) or "<tr><td colspan='2' style='color:#4CAF50;'>No subscales elevated ≥ 65</td></tr>"

    msg=MIMEMultipart('mixed')
    msg['From']=GMAIL_USER; msg['To']=RECIPIENT_EMAIL
    msg['Subject']=f"[Conners CPRS-R:L] {child_name} — {date_str}"
    body=f"""<html><body style="font-family:Georgia,serif;color:#1C1917;background:#F7F3EE;padding:20px;">
  <div style="max-width:560px;margin:0 auto;background:white;border:1px solid #DDD5C8;border-radius:4px;padding:28px;">
    <h2 style="font-weight:300;font-size:20px;color:#1C1917;margin-bottom:4px;">Conners' CPRS-R:L Report</h2>
    <p style="color:#8B7355;font-size:11px;margin-top:0;text-transform:uppercase;letter-spacing:.08em;">
      Conners' Parent Rating Scale — Revised: Long Version</p>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <table style="width:100%;font-size:13px;border-collapse:collapse;">
      <tr><td style="padding:5px 0;color:#8B7355;width:40%;">Child</td><td><strong>{child_name}</strong></td></tr>
      <tr><td style="padding:5px 0;color:#8B7355;">Date</td><td>{date_str}</td></tr>
    </table>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <p style="font-size:12px;color:#8B7355;font-weight:bold;margin-bottom:6px;">Elevated Subscales (T≥65)</p>
    <table style="width:100%;font-size:12px;border-collapse:collapse;">{elev_html}</table>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <p style="font-size:12px;line-height:1.6;">Two documents attached:<br>
    📄 <strong>English Report (PDF)</strong> — Full clinical report<br>
    📝 <strong>Arabic Report (Word)</strong> — التقرير السريري بالعربية</p>
    <p style="font-size:10px;color:#8B7355;font-style:italic;">Confidential — for the treating clinician only.</p>
  </div></body></html>"""
    msg.attach(MIMEText(body,'html'))
    # EN PDF
    buf_pdf_en.seek(0)
    part_pdf=MIMEBase('application','pdf')
    part_pdf.set_payload(buf_pdf_en.read()); encoders.encode_base64(part_pdf)
    part_pdf.add_header('Content-Disposition','attachment',filename=fn_pdf_en)
    msg.attach(part_pdf)
    # AR Word
    buf_word_ar.seek(0)
    part_doc=MIMEBase('application','vnd.openxmlformats-officedocument.wordprocessingml.document')
    part_doc.set_payload(buf_word_ar.read()); encoders.encode_base64(part_doc)
    part_doc.add_header('Content-Disposition','attachment',filename=fn_word_ar)
    msg.attach(part_doc)
    with smtplib.SMTP_SSL('smtp.gmail.com',465) as srv:
        srv.login(GMAIL_USER,GMAIL_PASS)
        srv.sendmail(GMAIL_USER,RECIPIENT_EMAIL,msg.as_string())

def send_email_en(child_name, buf_pdf_en, fn_pdf_en, scores):
    """English mode: attach English PDF only"""
    date_str=date.today().strftime('%B %d, %Y')
    elevated=[(k,scores[k]["t"]) for k in "ABCDEFGHIJKLMN" if scores[k]["t"]>=65]
    elev_html="".join(
        f"<tr><td style='padding:4px 0;color:#8B7355;'>{SUBSCALES[k]['name_en']}</td>"
        f"<td><strong style='color:#D9534F;'>T={t}</strong></td></tr>"
        for k,t in elevated
    ) or "<tr><td colspan='2' style='color:#4CAF50;'>No subscales elevated ≥ 65</td></tr>"

    msg=MIMEMultipart('mixed')
    msg['From']=GMAIL_USER; msg['To']=RECIPIENT_EMAIL
    msg['Subject']=f"[Conners CPRS-R:L] {child_name} — {date_str}"
    body=f"""<html><body style="font-family:Georgia,serif;color:#1C1917;background:#F7F3EE;padding:20px;">
  <div style="max-width:560px;margin:0 auto;background:white;border:1px solid #DDD5C8;border-radius:4px;padding:28px;">
    <h2 style="font-weight:300;font-size:20px;color:#1C1917;margin-bottom:4px;">Conners' CPRS-R:L Report</h2>
    <p style="color:#8B7355;font-size:11px;margin-top:0;text-transform:uppercase;letter-spacing:.08em;">
      Conners' Parent Rating Scale — Revised: Long Version</p>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <table style="width:100%;font-size:13px;border-collapse:collapse;">
      <tr><td style="padding:5px 0;color:#8B7355;width:40%;">Child</td><td><strong>{child_name}</strong></td></tr>
      <tr><td style="padding:5px 0;color:#8B7355;">Date</td><td>{date_str}</td></tr>
    </table>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <p style="font-size:12px;color:#8B7355;font-weight:bold;margin-bottom:6px;">Elevated Subscales (T≥65)</p>
    <table style="width:100%;font-size:12px;border-collapse:collapse;">{elev_html}</table>
    <hr style="border:none;border-top:1px solid #DDD5C8;margin:16px 0;">
    <p style="font-size:12px;line-height:1.6;">English clinical report attached as PDF.</p>
    <p style="font-size:10px;color:#8B7355;font-style:italic;">Confidential — for the treating clinician only.</p>
  </div></body></html>"""
    msg.attach(MIMEText(body,'html'))
    buf_pdf_en.seek(0)
    part=MIMEBase('application','pdf')
    part.set_payload(buf_pdf_en.read()); encoders.encode_base64(part)
    part.add_header('Content-Disposition','attachment',filename=fn_pdf_en)
    msg.attach(part)
    with smtplib.SMTP_SSL('smtp.gmail.com',465) as srv:
        srv.login(GMAIL_USER,GMAIL_PASS)
        srv.sendmail(GMAIL_USER,RECIPIENT_EMAIL,msg.as_string())

# ══════════════════════════════════════════════════════════════
#  PAGE CONFIG & MMPI THEME CSS
# ══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Conners CPRS-R:L | مقياس كونرز",
    page_icon="🧠",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@300;400;500&family=Jost:wght@300;400;500&display=swap');
:root{--cream:#F7F3EE;--deep:#1C1917;--warm:#8B7355;--accent:#C4956A;--border:#DDD5C8;--selected:#2D2926;}
#MainMenu{visibility:hidden!important;display:none!important;}
header[data-testid="stHeader"]{visibility:hidden!important;display:none!important;}
footer{visibility:hidden!important;display:none!important;}
[data-testid="stToolbar"],[data-testid="stDecoration"],[data-testid="stStatusWidget"],[data-testid="stActionButton"]{display:none!important;}
a[href*="streamlit.io"]{display:none!important;}
[class*="viewerBadge"],[class*="ProfileBadge"]{display:none!important;}
html,body,[data-theme="dark"],[data-theme="light"]{color-scheme:light only!important;}
[data-testid="stAppViewContainer"],.stApp{background-color:#F7F3EE!important;color:#1C1917!important;}
html,body,[class*="css"]{font-family:'Jost',sans-serif;background-color:var(--cream);color:var(--deep);}
.stApp{background-color:var(--cream);}

.page-header{text-align:center;padding:2.5rem 0 1.5rem;border-bottom:1px solid var(--border);margin-bottom:1.5rem;}
.page-header h1{font-family:'Cormorant Garamond',serif;font-size:2.2rem;font-weight:300;margin-bottom:.3rem;color:var(--deep);}
.page-header p{color:var(--warm);font-size:.82rem;letter-spacing:.05em;}

/* Language toggle buttons */
.lang-active > button{background:var(--selected)!important;color:var(--cream)!important;}

/* q-card — same as MMPI */
.q-card{background:white;border:1px solid var(--border);border-radius:4px;padding:1.2rem 1.5rem .5rem;margin-bottom:.8rem;}
.q-card:hover{border-color:var(--accent);}
.q-num{font-size:.68rem;font-weight:500;letter-spacing:.06em;color:var(--accent);margin-bottom:.3rem;}
.q-text{font-family:'Cormorant Garamond',serif;font-size:1.05rem;color:var(--deep);line-height:1.6;margin-bottom:.8rem;}

/* Radio pills — identical to MMPI */
div[data-testid="stRadio"]>label{display:none;}
div[data-testid="stRadio"]>div{gap:.4rem!important;flex-direction:row!important;flex-wrap:wrap!important;}
div[data-testid="stRadio"]>div>label{background:var(--cream)!important;border:1px solid var(--border)!important;border-radius:20px!important;padding:.4rem 1.2rem!important;cursor:pointer!important;font-size:.85rem!important;color:var(--deep)!important;font-family:'Jost',sans-serif!important;transition:all .15s ease!important;white-space:nowrap!important;}
div[data-testid="stRadio"]>div>label:hover{border-color:var(--accent)!important;background:#FDF9F4!important;}

/* RTL radio for Arabic */
.rtl-radio div[data-testid="stRadio"]>div{flex-direction:row-reverse!important;justify-content:flex-start!important;}

/* Progress bar */
.progress-wrap{background:var(--border);border-radius:2px;height:3px;margin:1rem 0 .5rem;}
.progress-fill{height:3px;border-radius:2px;background:linear-gradient(90deg,var(--warm),var(--accent));}

/* Submit button */
.stButton>button{background:var(--selected)!important;color:var(--cream)!important;border:none!important;padding:.75rem 2.5rem!important;font-family:'Jost',sans-serif!important;font-size:.82rem!important;font-weight:500!important;letter-spacing:.08em!important;border-radius:2px!important;transition:background .2s!important;}
.stButton>button:hover{background:var(--warm)!important;}

/* Thank-you screen */
.thank-you{text-align:center;padding:5rem 2rem;}
.thank-you h2{font-family:'Cormorant Garamond',serif;font-size:2.2rem;font-weight:300;margin-bottom:1rem;color:var(--deep);}
.thank-you p{color:var(--warm);font-size:.95rem;max-width:420px;margin:0 auto;line-height:1.9;}

/* Inputs */
div[data-testid="stTextInput"] input{background:white!important;border:1px solid var(--border)!important;border-radius:3px!important;font-family:'Jost',sans-serif!important;color:var(--deep)!important;}
div[data-testid="stSelectbox"] div{background:white!important;border:1px solid var(--border)!important;border-radius:3px!important;}
</style>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  SESSION STATE
# ══════════════════════════════════════════════════════════════
if "lang"         not in st.session_state: st.session_state.lang="en"
if "responses"    not in st.session_state: st.session_state.responses={}
if "submitted"    not in st.session_state: st.session_state.submitted=False
if "report_done"  not in st.session_state: st.session_state.report_done=False
if "access_granted" not in st.session_state: st.session_state.access_granted=False

# ══════════════════════════════════════════════════════════════
#  ACCESS CODE GATE
# ══════════════════════════════════════════════════════════════
if not st.session_state.access_granted:
    if os.path.exists(LOGO_FILE):
        c1,c2,c3=st.columns([1,2,1])
        with c2: st.image(LOGO_FILE, use_container_width=True)
    st.markdown("""<div class="page-header">
        <p>Confidential Clinical Assessment</p>
        <h1>Conners CPRS-R:L</h1>
    </div>""", unsafe_allow_html=True)
    st.markdown("""<div style="max-width:360px;margin:0 auto;padding:2rem 0;text-align:center;">
        <p style="color:#8B7355;font-size:.9rem;margin-bottom:1.5rem;line-height:1.8;">
            This assessment is available to referred patients only.<br>
            Please enter the access code provided by your clinician.
        </p>
    </div>""", unsafe_allow_html=True)
    col_a,col_b,col_c=st.columns([1,2,1])
    with col_b:
        code=st.text_input("Access code",type="password",
                           placeholder="Enter access code",label_visibility="collapsed")
        if st.button("Enter",use_container_width=True):
            valid_codes=[c.strip() for c in st.secrets.get("ACCESS_CODE","").split(",")]
            if code.strip() in valid_codes:
                st.session_state.access_granted=True; st.rerun()
            else:
                st.markdown("""<div style="background:#FFF0F0;border-left:3px solid #D9534F;
                    padding:.8rem 1rem;border-radius:0 4px 4px 0;
                    font-size:.88rem;color:#7A1A1A;margin:.5rem 0;">
                    &#9888; Incorrect access code. Please check and try again.
                </div>""", unsafe_allow_html=True)
    st.stop()

# ══════════════════════════════════════════════════════════════
#  RESULT DISPLAY (after submit) — email sent, thank-you only
# ══════════════════════════════════════════════════════════════
if st.session_state.report_done:
    scores     = st.session_state["scores"]
    child_name = st.session_state["child_name"]
    lang       = st.session_state.lang

    if os.path.exists(LOGO_FILE):
        c1,c2,c3=st.columns([1,2,1])
        with c2: st.image(LOGO_FILE, use_container_width=True)

    if lang=="en":
        st.markdown(f"""<div class="thank-you">
            <h2>Report Submitted Successfully</h2>
            <p style="font-size:1.05rem;margin-top:.5rem;">{child_name}</p>
            <p style="font-size:.9rem;color:#8B7355;margin-top:.8rem;">
                The report has been sent to the clinic email.<br>
                You may now close this window.
            </p>
        </div>""", unsafe_allow_html=True)
        _,btn_col,_=st.columns([2,2,2])
        with btn_col:
            if st.button("↺ New Assessment", use_container_width=True):
                for k in list(st.session_state.keys()):
                    if k not in {"lang","access_granted"}: del st.session_state[k]
                st.session_state.responses={}; st.session_state.submitted=False
                st.session_state.report_done=False; st.rerun()
    else:
        st.markdown(f"""<div class="thank-you" style="direction:rtl;">
            <h2>تم إرسال التقرير بنجاح</h2>
            <p style="font-size:1.05rem;margin-top:.5rem;">{child_name}</p>
            <p style="font-size:.9rem;color:#8B7355;margin-top:.8rem;">
                تم إرسال التقرير إلى البريد الإلكتروني للعيادة.<br>
                يمكنك إغلاق هذه النافذة الآن.
            </p>
        </div>""", unsafe_allow_html=True)
        _,btn_col,_=st.columns([2,2,2])
        with btn_col:
            if st.button("↺ تقييم جديد", use_container_width=True):
                for k in list(st.session_state.keys()):
                    if k not in {"lang","access_granted"}: del st.session_state[k]
                st.session_state.responses={}; st.session_state.submitted=False
                st.session_state.report_done=False; st.rerun()

    st.stop()

# ══════════════════════════════════════════════════════════════
#  FORM
# ══════════════════════════════════════════════════════════════
lang = st.session_state.lang

# Logo
if os.path.exists(LOGO_FILE):
    c1,c2,c3=st.columns([1,2,1])
    with c2: st.image(LOGO_FILE, use_container_width=True)

# Page header
if lang=="en":
    st.markdown("""<div class="page-header">
        <p>Confidential Clinical Assessment</p>
        <h1>Conners CPRS-R:L</h1>
        <p>Conners' Parent Rating Scale — Revised: Long Version</p>
    </div>""", unsafe_allow_html=True)
    st.markdown("""<div style="background:white;border:1px solid #DDD5C8;border-radius:6px;
        padding:1.4rem 1.8rem;margin:1.2rem 0 1.6rem;line-height:1.85;font-size:.9rem;color:#1C1917;">
        <div style="color:#8B7355;font-size:.72rem;font-weight:500;letter-spacing:.09em;
             text-transform:uppercase;margin-bottom:.7rem;">About this Assessment</div>
        <p style="margin:0 0 .8rem;">The <strong>Conners' Parent Rating Scale — Revised: Long Version (CPRS-R:L)</strong>
        is a standardised, evidence-based questionnaire widely used to evaluate behavioural, emotional,
        and attentional concerns in children aged 3–17 years — with a particular focus on symptoms
        associated with <strong>Attention-Deficit/Hyperactivity Disorder (ADHD)</strong>.</p>
        <p style="margin:0 0 .8rem;">You will be asked to rate <strong>80 statements</strong> describing
        your child's behaviour <strong>over the past month</strong>. Each statement is rated on a
        four-point scale from <em>Not at all</em> to <em>Very much</em>. There are no right or wrong answers —
        honest and accurate ratings lead to the most useful clinical picture.</p>
        <p style="margin:0;color:#8B7355;font-size:.84rem;">
        ⏱ Estimated time: 10–15 minutes &nbsp;·&nbsp;
        🔒 All responses are strictly confidential &nbsp;·&nbsp;
        📧 A report will be sent automatically to the clinician upon submission</p>
    </div>""", unsafe_allow_html=True)
else:
    st.markdown("""<div class="page-header" style="direction:rtl;">
        <p>تقييم سريري سري</p>
        <h1>مقياس كونرز للوالدين</h1>
        <p>CPRS-R:L — نسخة مراجعة طويلة</p>
    </div>""", unsafe_allow_html=True)
    st.markdown("""<div style="background:white;border:1px solid #DDD5C8;border-radius:6px;
        padding:1.4rem 1.8rem;margin:1.2rem 0 1.6rem;line-height:1.9;font-size:.9rem;color:#1C1917;
        direction:rtl;text-align:right;">
        <div style="color:#8B7355;font-size:.72rem;font-weight:500;letter-spacing:.05em;
             margin-bottom:.7rem;">عن هذا التقييم</div>
        <p style="margin:0 0 .8rem;"><strong>مقياس كونرز للوالدين — النسخة المراجعة الطويلة (CPRS-R:L)</strong>
        هو استبيان موحَّد ومعتمد بحثياً يُستخدم على نطاق واسع لتقييم المخاوف السلوكية والانفعالية
        والانتباهية لدى الأطفال من سن 3 إلى 17 سنة، مع تركيز خاص على أعراض
        <strong>اضطراب نقص الانتباه وفرط الحركة (ADHD)</strong>.</p>
        <p style="margin:0 0 .8rem;">ستُطلب منك تقييم <strong>80 عبارة</strong> تصف سلوك طفلك
        <strong>خلال الشهر الماضي</strong>. تُقيَّم كل عبارة على مقياس رباعي من
        <em>أبداً</em> إلى <em>كثيراً جداً</em>. لا توجد إجابات صحيحة أو خاطئة —
        التقييم الصادق والدقيق يُعطي الصورة السريرية الأكثر فائدة.</p>
        <p style="margin:0;color:#8B7355;font-size:.84rem;">
        ⏱ الوقت المتوقع: 10–15 دقيقة &nbsp;·&nbsp;
        🔒 جميع الإجابات سرية تماماً &nbsp;·&nbsp;
        📧 سيُرسل تقرير تلقائياً إلى الطبيب المعالج فور الإرسال</p>
    </div>""", unsafe_allow_html=True)

# Language toggle
c1,c2,c3=st.columns([2,2,4])
with c1:
    if st.button("English",use_container_width=True,
                 type="primary" if lang=="en" else "secondary"):
        st.session_state.lang="en"; st.session_state.responses={}
        st.session_state.submitted=False; st.session_state.report_done=False; st.rerun()
with c2:
    if st.button("العربية",use_container_width=True,
                 type="primary" if lang=="ar" else "secondary"):
        st.session_state.lang="ar"; st.session_state.responses={}
        st.session_state.submitted=False; st.session_state.report_done=False; st.rerun()

st.markdown("<br>", unsafe_allow_html=True)

# ── Child info fields ──
if lang=="en":
    st.markdown("""<div style="font-family:'Cormorant Garamond',serif;font-size:1.1rem;
        font-weight:400;color:#8B7355;letter-spacing:.08em;text-transform:uppercase;
        margin-bottom:.8rem;padding-bottom:.4rem;border-bottom:1px solid #DDD5C8;">
        Child Information</div>""", unsafe_allow_html=True)
    c1,c2,c3=st.columns(3)
    with c1:
        child_name=st.text_input("Child's Full Name",placeholder="First and last name",key="child_name_inp")
        child_age=st.selectbox("Age (years)",options=["—"]+[str(i) for i in range(3,18)],key="child_age_inp")
        child_age = "" if child_age=="—" else child_age
    with c2:
        child_gender=st.radio("Gender",["Male","Female"],key="child_gender_inp",
                              horizontal=True,label_visibility="visible")
        grade_opts=["—","KG","Grade 1","Grade 2","Grade 3","Grade 4","Grade 5",
                    "Grade 6","Grade 7","Grade 8","Grade 9","Grade 10","Grade 11","Grade 12"]
        child_grade=st.selectbox("School Grade",options=grade_opts,key="child_grade_inp")
        child_grade = "" if child_grade=="—" else child_grade
    with c3:
        rater=st.text_input("Rater's Name (Parent / Caregiver)",placeholder="Name",key="rater_inp")
        relationship=st.selectbox(
            "Relationship to Child",
            options=["—","Mother","Father","Grandmother","Grandfather","Guardian","Other"],
            key="rel_sel_en"
        )
        relationship = "" if relationship=="—" else relationship

    st.markdown(f"""<div style="background:white;border:1px solid #DDD5C8;border-radius:4px;
        padding:1rem 1.4rem;margin:1.2rem 0;font-size:.88rem;color:#1C1917;line-height:1.9;">
        <span style="color:#8B7355;font-weight:500;letter-spacing:.06em;text-transform:uppercase;font-size:.75rem;">Instructions</span><br>
        Below are common problems that children have. Please rate each item based on your child's behaviour
        <strong>in the last month</strong>. For each item, choose:<br>
        <strong>0</strong> — Not at all &nbsp;·&nbsp;
        <strong>1</strong> — Just a little &nbsp;·&nbsp;
        <strong>2</strong> — Pretty much &nbsp;·&nbsp;
        <strong>3</strong> — Very much
    </div>""", unsafe_allow_html=True)
    SCALE_OPTS=["0 — Not at all","1 — Just a little","2 — Pretty much","3 — Very much"]
    ITEMS=ITEMS_EN
    item_label="Item"

else:
    st.markdown("""<div style="font-family:'Cormorant Garamond',serif;font-size:1.1rem;
        font-weight:400;color:#8B7355;letter-spacing:.05em;
        margin-bottom:.8rem;padding-bottom:.4rem;border-bottom:1px solid #DDD5C8;
        direction:rtl;text-align:right;">
        بيانات الطفل</div>""", unsafe_allow_html=True)
    c1,c2,c3=st.columns(3)
    with c1:
        child_name=st.text_input("اسم الطفل (بالإنجليزية)",placeholder="e.g. Ahmed Hassan",key="child_name_inp")
        child_age=st.selectbox("السن (بالسنوات)",options=["—"]+[str(i) for i in range(3,18)],key="child_age_inp")
        child_age = "" if child_age=="—" else child_age
    with c2:
        child_gender=st.radio("النوع",["ذكر","أنثى"],key="child_gender_inp",
                              horizontal=True,label_visibility="visible")
        grade_opts_ar=["—","KG","الصف الأول","الصف الثاني","الصف الثالث","الصف الرابع",
                       "الصف الخامس","الصف السادس","الصف السابع","الصف الثامن",
                       "الصف التاسع","الصف العاشر","الصف الحادي عشر","الصف الثاني عشر"]
        child_grade=st.selectbox("الصف الدراسي",options=grade_opts_ar,key="child_grade_inp")
        child_grade = "" if child_grade=="—" else child_grade
    with c3:
        rater=st.text_input("اسم المُقيِّم (ولي الأمر)",placeholder="الاسم",key="rater_inp")
        relationship=st.selectbox(
            "صلة القرابة بالطفل",
            options=["—","الأم","الأب","الجدة","الجد","وصي","أخرى"],
            key="rel_sel_ar"
        )
        relationship = "" if relationship=="—" else relationship

    st.markdown(f"""<div style="background:white;border:1px solid #DDD5C8;border-radius:4px;
        padding:1rem 1.4rem;margin:1.2rem 0;font-size:.88rem;color:#1C1917;line-height:1.9;
        direction:rtl;text-align:right;">
        <span style="color:#8B7355;font-weight:500;letter-spacing:.04em;font-size:.75rem;">التعليمات</span><br>
        فيما يلي قائمة بالمشكلات الشائعة عند الأطفال. يرجى تقييم كل بند بناءً على سلوك طفلك
        <strong>خلال الشهر الماضي</strong>. اختر لكل بند:<br>
        <strong>0</strong> — أبداً &nbsp;·&nbsp;
        <strong>1</strong> — أحياناً &nbsp;·&nbsp;
        <strong>2</strong> — إلى حد ما &nbsp;·&nbsp;
        <strong>3</strong> — كثيراً جداً
    </div>""", unsafe_allow_html=True)
    SCALE_OPTS=["0 — أبداً","1 — أحياناً","2 — إلى حد ما","3 — كثيراً جداً"]
    ITEMS=ITEMS_AR
    item_label="بند"

# ══════════════════════════════════════════════════════════════
#  80 ITEMS — MMPI q-card style
# ══════════════════════════════════════════════════════════════
responses=st.session_state.responses
all_answered=True
direction='rtl' if lang=='ar' else 'ltr'
align='right' if lang=='ar' else 'left'

for idx, item_text in enumerate(ITEMS):
    item_num=idx+1
    st.markdown(f"""<div class="q-card" style="direction:{direction};">
        <div class="q-num" style="text-align:{align};">{item_label} {item_num} / 80</div>
        <div class="q-text" style="text-align:{align};">{item_text}</div>
    </div>""", unsafe_allow_html=True)

    saved=responses.get(item_num)
    choice=st.radio(
        f"item_{item_num}",
        SCALE_OPTS,
        index=saved,
        key=f"resp_{item_num}",
        horizontal=True,
        label_visibility="collapsed"
    )
    if choice is None:
        all_answered=False
    else:
        val=int(choice[0])
        responses[item_num]=val
        st.session_state.responses[item_num]=val

# ══════════════════════════════════════════════════════════════
#  PROGRESS & SUBMIT
# ══════════════════════════════════════════════════════════════
answered_count=len([v for v in responses.values() if v is not None])
pct=int((answered_count/80)*100)
prog_text=f"{answered_count} of 80 answered" if lang=="en" else f"{answered_count} من 80 بنداً"

st.markdown(f"""
<div style="text-align:center;font-size:.78rem;color:#8B7355;
            letter-spacing:.06em;margin-top:1.5rem;">{prog_text}</div>
<div class="progress-wrap">
    <div class="progress-fill" style="width:{pct}%"></div>
</div>""", unsafe_allow_html=True)

if not all_answered and answered_count>0:
    warn="⚠ Please answer all 80 items before submitting." if lang=="en" \
         else "⚠ يرجى الإجابة على جميع البنود الـ 80 قبل الإرسال."
    st.markdown(f"""<div style="background:#FFF8F0;border-left:3px solid #E07B39;
        padding:1rem 1.2rem;border-radius:0 4px 4px 0;
        font-size:.88rem;color:#7A3D1A;margin:1rem 0;">{warn}</div>""",
        unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)
btn_label="✦ Generate Report" if lang=="en" else "✦ توليد التقرير"
col_btn,_=st.columns([2,3])
with col_btn:
    submit=st.button(btn_label,use_container_width=True,disabled=(answered_count<80))

if submit and answered_count==80:
    child_name_v=child_name  or ("Child" if lang=="en" else "الطفل")
    child_age_v =child_age   or "—"
    rater_v     =rater       or ("Parent" if lang=="en" else "ولي الأمر")
    gender_v    =child_gender
    # For Arabic UI: translate gender to English for the English PDF
    gender_v_en = gender_v
    if lang == "ar":
        gender_v_en = "Male" if gender_v == "ذكر" else "Female"
    responses_v =dict(st.session_state.responses)

    spinner_txt=("⏳ Scoring and generating report..." if lang=="en"
                 else "⏳ جاري الحساب وإنشاء التقارير...")
    with st.spinner(spinner_txt):
        scores   =compute_scores(st.session_state.responses)
        report_en=generate_report_en(child_name_v,child_age_v,gender_v_en,rater_v,scores)

        # Arabic mode → also generate Arabic report
        report_ar=""
        if lang=="ar":
            report_ar=generate_report_ar(child_name_v,child_age_v,gender_v,rater_v,scores)

        # Auto-email
        try:
            bar_b=make_bar_chart(scores,lang); pie_b=make_pie_chart(responses_v)
            # English report → always PDF
            buf_pdf_en_=build_pdf_report_en(report_en,scores,bar_b,pie_b,
                                             child_name_v,child_age_v,gender_v_en,rater_v,responses_v)
            fn_pdf_en_=f"{child_name_v.replace(' ','_')}_Conners_EN.pdf"
            if lang=="ar":
                # Arabic report → Word
                buf_word_ar_=build_word_report(report_ar,scores,bar_b,pie_b,
                                               child_name_v,child_age_v,gender_v,rater_v,"ar",responses_v)
                fn_word_ar_=f"{child_name_v.replace(' ','_')}_Conners_AR.docx"
                send_email_ar(child_name_v,buf_pdf_en_,buf_word_ar_,fn_pdf_en_,fn_word_ar_,scores)
            else:
                send_email_en(child_name_v,buf_pdf_en_,fn_pdf_en_,scores)
        except Exception:
            pass  # email failure is silent

        st.session_state["scores"]      =scores
        st.session_state["report_en"]   =report_en
        st.session_state["report_ar"]   =report_ar
        st.session_state["child_name"]  =child_name_v
        st.session_state["child_age"]   =child_age_v
        st.session_state["child_gender"]=gender_v_en
        st.session_state["rater"]       =rater_v
        st.session_state.report_done    =True
        st.rerun()
